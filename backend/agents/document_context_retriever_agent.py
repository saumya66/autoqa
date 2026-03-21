"""
Document Context Retriever Agent

Processes documents (PRDs, specifications, text) and extracts
structured requirements for test generation.
"""

import io
from typing import Any, Optional

from .base_agent import BaseAgent


class DocumentContextRetrieverAgent(BaseAgent):
    """
    Agent for extracting requirements from documents.
    
    Can process:
    - PRD documents (PDF, DOCX)
    - Plain text specifications
    - Feature descriptions
    - Acceptance criteria
    """
    
    @property
    def system_prompt(self) -> str:
        return """You are a QA requirements analyst. Your task is to analyze product documentation and extract structured test requirements.

When analyzing a document, extract:

1. FEATURE OVERVIEW
   - Feature name
   - Brief summary
   - Target users (if mentioned)

2. REQUIREMENTS - Each individual requirement should be:
   - Clearly stated
   - Categorized (functional, validation, error_handling, edge_case, ux)
   - Prioritized (must, should, could, wont)

3. USER FLOWS - Step-by-step processes:
   - Flow name
   - Ordered steps
   - Expected outcomes

4. EDGE CASES & ERROR HANDLING:
   - What can go wrong
   - Expected error messages
   - Recovery steps

5. ACCEPTANCE CRITERIA:
   - Measurable conditions for success

Respond ONLY with valid JSON in this exact format:
{
  "feature_name": "Add to Bag",
  "summary": "Allows users to add products to their shopping bag from product details page",
  "target_users": ["shoppers", "guests", "registered_users"],
  "requirements": [
    {
      "id": "REQ-001",
      "text": "User must be able to add a product to bag from PDP",
      "category": "functional",
      "priority": "must"
    },
    {
      "id": "REQ-002",
      "text": "Size must be selected before adding to bag",
      "category": "validation",
      "priority": "must"
    },
    {
      "id": "REQ-003",
      "text": "Show error toast if product goes out of stock",
      "category": "error_handling",
      "priority": "must"
    },
    {
      "id": "REQ-004",
      "text": "Show quantity selector if same product added again",
      "category": "functional",
      "priority": "should"
    }
  ],
  "user_flows": [
    {
      "name": "Happy Path - Add to Bag",
      "steps": [
        "User navigates to product details page",
        "User selects size",
        "User taps Add to Bag button",
        "Success toast appears",
        "Bag icon updates with count"
      ],
      "expected_outcome": "Product is in bag with correct size"
    }
  ],
  "edge_cases": [
    "Product goes out of stock while user is on PDP",
    "Network error during add to bag API call",
    "User tries to add more than available stock"
  ],
  "acceptance_criteria": [
    "Product appears in bag with correct details",
    "Bag count increments correctly",
    "Success toast shows for 3 seconds",
    "Can continue shopping after adding"
  ]
}

Extract as much structured information as possible. If something is not mentioned in the document, omit that field."""

    def parse_response(self, response_text: str) -> Any:
        """Parse the JSON response from Gemini."""
        return self.extract_json(response_text)
    
    def extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extract text content from a PDF file.
        
        Args:
            pdf_bytes: The PDF file as bytes.
        
        Returns:
            Extracted text content.
        """
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            text_content = []
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
            
            doc.close()
            return "\n\n".join(text_content)
        except ImportError:
            raise ImportError("PyMuPDF (fitz) is required for PDF processing. Install with: pip install PyMuPDF")
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {e}")
    
    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """
        Extract text content from a DOCX file.
        
        Args:
            docx_bytes: The DOCX file as bytes.
        
        Returns:
            Extracted text content.
        """
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(docx_bytes))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            return "\n\n".join(paragraphs)
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {e}")
    
    def process_text(self, text: str, additional_context: str = "") -> Optional[dict]:
        """
        Process plain text and extract requirements.
        
        Args:
            text: The document text content.
            additional_context: Optional additional context.
        
        Returns:
            Structured requirements dict or None if processing failed.
        """
        prompt = f"""Analyze this document and extract structured requirements:

--- DOCUMENT START ---
{text[:15000]}  # Limit to ~15k chars to avoid token limits
--- DOCUMENT END ---"""
        
        if additional_context:
            prompt += f"\n\nAdditional context: {additional_context}"
        
        try:
            response = self.call_llm(
                user_prompt=prompt,
                max_tokens=4096
            )
            return self.parse_response(response)
        except Exception as e:
            print(f"[DocumentContextRetrieverAgent] Error processing text: {e}")
            return None
    
    def process_pdf(self, pdf_bytes: bytes, additional_context: str = "") -> Optional[dict]:
        """
        Process a PDF document and extract requirements.
        
        Args:
            pdf_bytes: The PDF file as bytes.
            additional_context: Optional additional context.
        
        Returns:
            Structured requirements dict or None if processing failed.
        """
        text = self.extract_text_from_pdf(pdf_bytes)
        return self.process_text(text, additional_context)
    
    def process_docx(self, docx_bytes: bytes, additional_context: str = "") -> Optional[dict]:
        """
        Process a DOCX document and extract requirements.
        
        Args:
            docx_bytes: The DOCX file as bytes.
            additional_context: Optional additional context.
        
        Returns:
            Structured requirements dict or None if processing failed.
        """
        text = self.extract_text_from_docx(docx_bytes)
        return self.process_text(text, additional_context)
    
    def process(
        self,
        content: bytes,
        file_type: str,
        additional_context: str = ""
    ) -> Optional[dict]:
        """
        Process a document of any supported type.
        
        Args:
            content: The file content as bytes.
            file_type: One of "pdf", "docx", "txt".
            additional_context: Optional additional context.
        
        Returns:
            Structured requirements dict or None if processing failed.
        """
        if file_type == "pdf":
            return self.process_pdf(content, additional_context)
        elif file_type == "docx":
            return self.process_docx(content, additional_context)
        elif file_type == "txt":
            text = content.decode("utf-8", errors="ignore")
            return self.process_text(text, additional_context)
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Use pdf, docx, or txt.")
